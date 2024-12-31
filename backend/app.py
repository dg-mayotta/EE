from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import openai
import re
import os
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

# Configuration
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
MAX_CHUNK_SIZE = 4000
MODEL_NAME = "gpt-4"
TEMPERATURE = 0.7

def process_docx(file):
    """
    Process a DOCX file and extract clean text while preserving structure
    """
    try:
        doc = Document(file)
        processed_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Clean the text while preserving important formatting
                clean_text = re.sub(r'\s+', ' ', paragraph.text)
                clean_text = re.sub(r'[^\w\s.,;:?¿!¡()\-\'\"]+', '', clean_text)
                
                # Add formatting indicators if paragraph has special styling
                if paragraph.style.name.startswith('Heading'):
                    clean_text = f"## {clean_text}"
                
                processed_text.append(clean_text.strip())

        # Process tables
        for table in doc.tables:
            processed_text.append("\nTable:")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                processed_text.append(" | ".join(row_text))

        return "\n\n".join(processed_text)
    except Exception as e:
        raise Exception(f"Error processing DOCX file: {str(e)}")

def chunk_text(text):
    """
    Split text into chunks that don't exceed GPT-4's token limit
    """
    # Approximate tokens (rough estimate: 4 chars = 1 token)
    max_chars = MAX_CHUNK_SIZE * 4
    
    # If text is short enough, return as single chunk
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    for paragraph in paragraphs:
        # If a single paragraph is too long, split it by sentences
        if len(paragraph) > max_chars:
            sentences = paragraph.split('. ')
            for sentence in sentences:
                if current_length + len(sentence) > max_chars:
                    # Save current chunk and start new one
                    chunks.append('\n\n'.join(current_chunk))
                    current_chunk = [sentence]
                    current_length = len(sentence)
                else:
                    current_chunk.append(sentence)
                    current_length += len(sentence)
        else:
            if current_length + len(paragraph) > max_chars:
                # Save current chunk and start new one
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [paragraph]
                current_length = len(paragraph)
            else:
                current_chunk.append(paragraph)
                current_length += len(paragraph)
    
    # Add the last chunk if it exists
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    return chunks

def analyze_with_gpt(text):
    """
    Analyze text using GPT-4
    """
    try:
        openai.api_key = OPENAI_API_KEY
        
        prompt = f"""
        Please analyze the following Spanish text and:
        1. Display the content in English in a clear way, preserving all the information
        2. Perform a comprehensive analysis
        
        Text to analyze:
        {text}
        """
        
        response = openai.ChatCompletion.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "You are an expert analyst. Provide clear, structured analysis of documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=TEMPERATURE,
            max_tokens=MAX_CHUNK_SIZE
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        raise Exception(f"Error in GPT analysis: {str(e)}")

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy"}), 200

@app.route('/analyze', methods=['POST'])
def analyze_document():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        if not file.filename.endswith('.docx'):
            return jsonify({"error": "Only DOCX files are supported"}), 400

        # Extract text from DOCX
        extracted_text = process_docx(file)
        if not extracted_text:
            return jsonify({"error": "Failed to extract text from document"}), 400
        
        # Split into chunks if necessary
        text_chunks = chunk_text(extracted_text)
        
        # Analyze with GPT-4
        results = []
        for chunk in text_chunks:
            analysis = analyze_with_gpt(chunk)
            results.append(analysis)
        
        # Combine results
        final_result = {
            "original_text": extracted_text,
            "analysis": results
        }
        
        return jsonify(final_result), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
