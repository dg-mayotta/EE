from docx import Document
import re

def process_docx(file):
    """
    Process a DOCX file and extract clean text while preserving structure
    """
    try:
        doc = Document(file)
        processed_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Clean the text while preserving important formatting
                clean_text = clean_paragraph_text(paragraph.text)
                
                # Add formatting indicators if paragraph has special styling
                if paragraph.style.name.startswith('Heading'):
                    clean_text = f"## {clean_text}"
                
                processed_text.append(clean_text)

        # Process tables
        for table in doc.tables:
            processed_text.append("\nTable:")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                processed_text.append(" | ".join(row_text))

        return "\n\n".join(processed_text)
    except Exception as e:
        raise Exception(f"Error processing DOCX file: {str(e)}")

def clean_paragraph_text(text):
    """
    Clean paragraph text while preserving important information
    """
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,;:?¿!¡()\-\'\"]+', '', text)
    
    return text.strip()
